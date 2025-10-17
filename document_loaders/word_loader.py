import os
import hashlib
from typing import List, Dict, Any, Optional
from .base import BaseDocumentLoader, Document
from docx import Document as DocxDocument


class WordFileLoader(BaseDocumentLoader):
    """
    Word文件加载器，用于加载.docx文件并分块处理
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        """
        初始化Word文件加载器
        
        Args:
            chunk_size: 分块大小
            chunk_overlap: 块重叠大小
        """
        super().__init__()
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def load(self, source: str, **kwargs) -> List[Document]:
        """
        加载Word文件
        
        Args:
            source: Word文件路径
            **kwargs: 其他可选参数
            
        Returns:
            加载的文档列表
            
        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 文件类型错误
        """
        # 检查文件是否存在
        if not os.path.exists(source):
            raise FileNotFoundError(f"文件不存在: {source}")
        
        # 检查文件类型
        if not source.lower().endswith('.docx'):
            raise ValueError(f"不支持的文件类型，需要.docx文件: {source}")
        
        # 设置分块参数（如果提供）
        chunk_size = kwargs.get('chunk_size', self.chunk_size)
        chunk_overlap = kwargs.get('chunk_overlap', self.chunk_overlap)
        
        # 临时保存原始分块参数
        original_chunk_size = self.chunk_size
        original_chunk_overlap = self.chunk_overlap
        
        # 更新分块参数
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        try:
            # 打开并读取Word文件
            doc = DocxDocument(source)
            text = ""
            
            # 提取所有段落的文本
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # 生成文档ID
            doc_id = hashlib.md5(source.encode()).hexdigest()
            
            # 分块文本
            chunks = self.chunk_text(text)
            
            # 创建文档列表
            documents = []
            for i, chunk in enumerate(chunks):
                # 为每个块生成唯一ID
                chunk_id = f"{doc_id}_chunk_{i}"
                
                # 创建文档对象
                document = Document(
                    id=chunk_id,
                    content=chunk,
                    metadata={
                        "source": source,
                        "file_name": os.path.basename(source),
                        "file_type": "docx",
                        "paragraph_count": len(doc.paragraphs),
                        "chunk_index": i,
                        "total_chunks": len(chunks)
                    },
                    chunk_id=chunk_id
                )
                documents.append(document)
            
            return documents
        
        finally:
            # 恢复原始分块参数
            self.chunk_size = original_chunk_size
            self.chunk_overlap = original_chunk_overlap
    
    @staticmethod
    def get_loader_info() -> Dict[str, Any]:
        """
        获取加载器信息
        
        Returns:
            加载器信息字典
        """
        return {
            "name": "WordFileLoader",
            "supported_extensions": [".docx"]
        }